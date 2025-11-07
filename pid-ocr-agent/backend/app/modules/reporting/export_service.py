"""
Document Export Service
Exports HAZOP and Instrument Index to various formats (PDF, Excel, Word)
"""
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from typing import Dict, List
import logging
from pathlib import Path
from datetime import datetime
import csv

logger = logging.getLogger(__name__)


class HAZOPExporter:
    """Export HAZOP studies to various formats"""

    def __init__(self):
        self.styles = getSampleStyleSheet()

    def export_to_pdf(self, hazop_data: Dict, output_path: str) -> str:
        """
        Export HAZOP study to PDF

        Args:
            hazop_data: HAZOP study data
            output_path: Output file path

        Returns:
            Path to generated PDF
        """
        try:
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            story = []

            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=self.styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1f4788'),
                spaceAfter=30,
                alignment=TA_CENTER
            )

            title = Paragraph("HAZOP Study Report", title_style)
            story.append(title)
            story.append(Spacer(1, 0.2 * inch))

            # Project information
            project_info = [
                ["Project Name:", hazop_data.get("project_name", "")],
                ["P&ID Reference:", hazop_data.get("pid_reference", "")],
                ["Study Date:", hazop_data.get("study_date", "")],
                ["Revision:", hazop_data.get("revision", "")],
                ["Status:", hazop_data.get("status", "")]
            ]

            info_table = Table(project_info, colWidths=[2 * inch, 4 * inch])
            info_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.grey),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))

            story.append(info_table)
            story.append(Spacer(1, 0.3 * inch))

            # Process each node
            for node in hazop_data.get("nodes", []):
                # Node header
                node_header = Paragraph(
                    f"<b>Node {node.get('node_number')}: {node.get('node_description')}</b>",
                    self.styles['Heading2']
                )
                story.append(node_header)
                story.append(Spacer(1, 0.1 * inch))

                # Design intent
                design_intent = Paragraph(
                    f"<b>Design Intent:</b> {node.get('design_intent', '')}",
                    self.styles['Normal']
                )
                story.append(design_intent)
                story.append(Spacer(1, 0.2 * inch))

                # Deviations table
                deviation_data = [["Guide Word", "Parameter", "Causes", "Consequences", "Safeguards", "Risk", "Recommendations"]]

                for dev in node.get("deviations", []):
                    deviation_data.append([
                        dev.get("guide_word", ""),
                        dev.get("parameter", ""),
                        "\n".join(dev.get("possible_causes", [])[:2]),
                        "\n".join(dev.get("consequences", [])[:2]),
                        "\n".join(dev.get("safeguards", [])[:2]),
                        str(dev.get("risk_ranking", "")),
                        "\n".join(dev.get("recommendations", [])[:1])
                    ])

                deviation_table = Table(
                    deviation_data,
                    colWidths=[0.8 * inch, 0.8 * inch, 1.5 * inch, 1.5 * inch, 1.3 * inch, 0.4 * inch, 1.2 * inch]
                )

                deviation_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP')
                ]))

                story.append(deviation_table)
                story.append(PageBreak())

            # Build PDF
            doc.build(story)

            logger.info(f"HAZOP PDF exported to {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error exporting HAZOP to PDF: {str(e)}")
            raise

    def export_to_word(self, hazop_data: Dict, output_path: str) -> str:
        """
        Export HAZOP study to Word document

        Args:
            hazop_data: HAZOP study data
            output_path: Output file path

        Returns:
            Path to generated Word document
        """
        try:
            doc = Document()

            # Title
            title = doc.add_heading("HAZOP Study Report", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Project information
            doc.add_heading("Project Information", 2)
            table = doc.add_table(rows=5, cols=2)
            table.style = 'Light Grid Accent 1'

            info_data = [
                ("Project Name:", hazop_data.get("project_name", "")),
                ("P&ID Reference:", hazop_data.get("pid_reference", "")),
                ("Study Date:", hazop_data.get("study_date", "")),
                ("Revision:", hazop_data.get("revision", "")),
                ("Status:", hazop_data.get("status", ""))
            ]

            for idx, (label, value) in enumerate(info_data):
                table.rows[idx].cells[0].text = label
                table.rows[idx].cells[1].text = str(value)

            doc.add_page_break()

            # Process each node
            for node in hazop_data.get("nodes", []):
                doc.add_heading(f"Node {node.get('node_number')}: {node.get('node_description')}", 2)

                doc.add_paragraph(f"Design Intent: {node.get('design_intent', '')}")

                # Deviations table
                deviation_table = doc.add_table(rows=1, cols=7)
                deviation_table.style = 'Light Grid Accent 1'

                headers = ["Guide Word", "Parameter", "Causes", "Consequences", "Safeguards", "Risk", "Recommendations"]
                header_cells = deviation_table.rows[0].cells
                for idx, header in enumerate(headers):
                    header_cells[idx].text = header

                for dev in node.get("deviations", []):
                    row_cells = deviation_table.add_row().cells
                    row_cells[0].text = dev.get("guide_word", "")
                    row_cells[1].text = dev.get("parameter", "")
                    row_cells[2].text = "\n".join(dev.get("possible_causes", []))
                    row_cells[3].text = "\n".join(dev.get("consequences", []))
                    row_cells[4].text = "\n".join(dev.get("safeguards", []))
                    row_cells[5].text = str(dev.get("risk_ranking", ""))
                    row_cells[6].text = "\n".join(dev.get("recommendations", []))

                doc.add_page_break()

            # Save document
            doc.save(output_path)

            logger.info(f"HAZOP Word document exported to {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error exporting HAZOP to Word: {str(e)}")
            raise

    def export_to_excel(self, hazop_data: Dict, output_path: str) -> str:
        """
        Export HAZOP study to Excel

        Args:
            hazop_data: HAZOP study data
            output_path: Output file path

        Returns:
            Path to generated Excel file
        """
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "HAZOP Study"

            # Styling
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=11)
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            # Project information
            ws['A1'] = "HAZOP Study Report"
            ws['A1'].font = Font(bold=True, size=16)
            ws.merge_cells('A1:G1')

            info_data = [
                ("Project Name:", hazop_data.get("project_name", "")),
                ("P&ID Reference:", hazop_data.get("pid_reference", "")),
                ("Study Date:", hazop_data.get("study_date", "")),
                ("Revision:", hazop_data.get("revision", ""))
            ]

            row = 3
            for label, value in info_data:
                ws.cell(row, 1, label).font = Font(bold=True)
                ws.cell(row, 2, value)
                row += 1

            row += 2

            # Process each node
            for node in hazop_data.get("nodes", []):
                # Node header
                ws.cell(row, 1, f"Node {node.get('node_number')}: {node.get('node_description')}").font = Font(bold=True, size=12)
                ws.merge_cells(f'A{row}:G{row}')
                row += 1

                ws.cell(row, 1, f"Design Intent: {node.get('design_intent', '')}")
                ws.merge_cells(f'A{row}:G{row}')
                row += 2

                # Headers
                headers = ["Guide Word", "Parameter", "Causes", "Consequences", "Safeguards", "Risk", "Recommendations"]
                for col, header in enumerate(headers, start=1):
                    cell = ws.cell(row, col, header)
                    cell.fill = header_fill
                    cell.font = header_font
                    cell.border = border
                    cell.alignment = Alignment(horizontal='center', vertical='center')

                row += 1

                # Deviations
                for dev in node.get("deviations", []):
                    ws.cell(row, 1, dev.get("guide_word", "")).border = border
                    ws.cell(row, 2, dev.get("parameter", "")).border = border
                    ws.cell(row, 3, "\n".join(dev.get("possible_causes", []))).border = border
                    ws.cell(row, 4, "\n".join(dev.get("consequences", []))).border = border
                    ws.cell(row, 5, "\n".join(dev.get("safeguards", []))).border = border
                    ws.cell(row, 6, dev.get("risk_ranking", "")).border = border
                    ws.cell(row, 7, "\n".join(dev.get("recommendations", []))).border = border
                    row += 1

                row += 2

            # Adjust column widths
            ws.column_dimensions['A'].width = 15
            ws.column_dimensions['B'].width = 15
            ws.column_dimensions['C'].width = 30
            ws.column_dimensions['D'].width = 30
            ws.column_dimensions['E'].width = 25
            ws.column_dimensions['F'].width = 8
            ws.column_dimensions['G'].width = 30

            # Save workbook
            wb.save(output_path)

            logger.info(f"HAZOP Excel file exported to {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error exporting HAZOP to Excel: {str(e)}")
            raise


class InstrumentIndexExporter:
    """Export Instrument Index to various formats"""

    def export_to_excel(self, index_data: Dict, output_path: str) -> str:
        """
        Export Instrument Index to Excel

        Args:
            index_data: Instrument index data
            output_path: Output file path

        Returns:
            Path to generated Excel file
        """
        try:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Instrument Index"

            # Styling
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=11)
            border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )

            # Title
            ws['A1'] = "Instrument Index"
            ws['A1'].font = Font(bold=True, size=16)
            ws.merge_cells('A1:H1')

            # Project info
            ws['A2'] = f"Project: {index_data.get('project_name', '')}"
            ws['A3'] = f"P&ID Reference: {index_data.get('pid_reference', '')}"
            ws['A4'] = f"Date: {index_data.get('date', '')}"

            # Headers
            headers = [
                "Tag", "Measured Variable", "Function", "Service Description",
                "Location", "P&ID Ref", "Range", "Signal Type"
            ]

            row = 6
            for col, header in enumerate(headers, start=1):
                cell = ws.cell(row, col, header)
                cell.fill = header_fill
                cell.font = header_font
                cell.border = border
                cell.alignment = Alignment(horizontal='center', vertical='center')

            row += 1

            # Instruments
            for inst in index_data.get("instruments", []):
                ws.cell(row, 1, inst.get("tag", "")).border = border
                ws.cell(row, 2, inst.get("measured_variable", "")).border = border
                ws.cell(row, 3, inst.get("function", "")).border = border
                ws.cell(row, 4, inst.get("service_description", "")).border = border
                ws.cell(row, 5, inst.get("location", "")).border = border
                ws.cell(row, 6, inst.get("pid_reference", "")).border = border

                range_str = ""
                if inst.get("range_min") and inst.get("range_max"):
                    range_str = f"{inst['range_min']}-{inst['range_max']} {inst.get('units', '')}"
                ws.cell(row, 7, range_str).border = border

                ws.cell(row, 8, inst.get("signal_type", "")).border = border
                row += 1

            # Adjust column widths
            ws.column_dimensions['A'].width = 15
            ws.column_dimensions['B'].width = 18
            ws.column_dimensions['C'].width = 15
            ws.column_dimensions['D'].width = 35
            ws.column_dimensions['E'].width = 15
            ws.column_dimensions['F'].width = 15
            ws.column_dimensions['G'].width = 20
            ws.column_dimensions['H'].width = 15

            # Save workbook
            wb.save(output_path)

            logger.info(f"Instrument Index Excel file exported to {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error exporting Instrument Index to Excel: {str(e)}")
            raise

    def export_to_csv(self, index_data: Dict, output_path: str) -> str:
        """
        Export Instrument Index to CSV

        Args:
            index_data: Instrument index data
            output_path: Output file path

        Returns:
            Path to generated CSV file
        """
        try:
            with open(output_path, 'w', newline='') as csvfile:
                fieldnames = [
                    "Tag", "Measured Variable", "Function", "Service Description",
                    "Location", "P&ID Reference", "Range Min", "Range Max",
                    "Units", "Signal Type", "Status"
                ]

                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                writer.writeheader()

                for inst in index_data.get("instruments", []):
                    writer.writerow({
                        "Tag": inst.get("tag", ""),
                        "Measured Variable": inst.get("measured_variable", ""),
                        "Function": inst.get("function", ""),
                        "Service Description": inst.get("service_description", ""),
                        "Location": inst.get("location", ""),
                        "P&ID Reference": inst.get("pid_reference", ""),
                        "Range Min": inst.get("range_min", ""),
                        "Range Max": inst.get("range_max", ""),
                        "Units": inst.get("units", ""),
                        "Signal Type": inst.get("signal_type", ""),
                        "Status": inst.get("status", "")
                    })

            logger.info(f"Instrument Index CSV file exported to {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Error exporting Instrument Index to CSV: {str(e)}")
            raise
